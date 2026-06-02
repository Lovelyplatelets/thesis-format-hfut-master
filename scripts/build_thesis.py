#!/usr/bin/env python3
"""HFUT master's dissertation DOCX generator.

Usage:
    python build_thesis.py input.md output.docx
"""

import re
import sys
from pathlib import Path

try:
    import yaml
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:
    print("Missing dependency. Install with: python -m pip install python-docx pyyaml")
    print(exc)
    sys.exit(1)


FONT_SIZE = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_EN = "Times New Roman"


def size_pt(size):
    if isinstance(size, (int, float)):
        return float(size)
    return FONT_SIZE[size]


def set_run_font(run, cn_font=FONT_SONG, en_font=FONT_EN, size="小四", bold=False, italic=False):
    run.font.name = en_font
    run.font.size = Pt(size_pt(size))
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def set_line_spacing_exact(paragraph, pt):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(int(pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")


def set_single_line(paragraph):
    paragraph.paragraph_format.line_spacing = 1


def set_spacing(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def set_first_line_indent(paragraph, chars=2, size="小四"):
    paragraph.paragraph_format.first_line_indent = Pt(size_pt(size) * chars)


def set_hanging_indent(paragraph, chars=2.5, size="五号"):
    width = Pt(size_pt(size) * chars)
    paragraph.paragraph_format.left_indent = width
    paragraph.paragraph_format.first_line_indent = -width


def add_text_paragraph(
    doc,
    text,
    cn_font=FONT_SONG,
    en_font=FONT_EN,
    size="小四",
    bold=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    first_indent=0,
    exact_line=20,
    before=0,
    after=0,
    style=None,
):
    p = doc.add_paragraph()
    if style:
        p.style = style
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold=bold)
    if first_indent:
        set_first_line_indent(p, first_indent, size)
    if exact_line:
        set_line_spacing_exact(p, exact_line)
    else:
        set_single_line(p)
    set_spacing(p, before, after)
    return p


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def set_section_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def unlink_headers_footers(section):
    for part in (
        section.header,
        section.first_page_header,
        section.even_page_header,
        section.footer,
        section.first_page_footer,
        section.even_page_footer,
    ):
        part.is_linked_to_previous = False


def set_footer_page_number(section):
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(p, "PAGE")
        for run in p.runs:
            set_run_font(run, FONT_SONG, FONT_EN, "五号")


def set_body_headers(section):
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False

    odd = section.header.paragraphs[0]
    clear_paragraph(odd)
    odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(odd, 'STYLEREF "Heading 1"')
    for run in odd.runs:
        set_run_font(run, FONT_SONG, FONT_EN, "五号")

    even = section.even_page_header.paragraphs[0]
    clear_paragraph(even)
    even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = even.add_run("合肥工业大学硕士学位论文")
    set_run_font(run, FONT_SONG, FONT_EN, "五号")


def parse_source(path):
    text = Path(path).read_text(encoding="utf-8")
    if text.startswith("---"):
        _, frontmatter, body = text.split("---", 2)
        meta = yaml.safe_load(frontmatter) or {}
    else:
        meta = {}
        body = text
    return meta, body.strip()


def parse_markdown_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"[-:\s]+", cell or "") for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def parse_body(body):
    lines = body.splitlines()
    elements = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            elements.append(("chapter", stripped[2:].strip()))
            i += 1
            continue
        if stripped.startswith("## "):
            elements.append(("section", stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith("### "):
            elements.append(("subsection", stripped[4:].strip()))
            i += 1
            continue

        fig_match = re.match(r'^!\[([^\]]*)\]\(([^)\s]+)(?:\s+"([^"]+)")?\)$', stripped)
        if fig_match:
            caption_zh = fig_match.group(1).strip()
            image_path = fig_match.group(2).strip()
            caption_en = (fig_match.group(3) or "").strip()
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("Fig "):
                caption_en = lines[i + 1].strip()
                i += 1
            elements.append(("figure", {"caption_zh": caption_zh, "caption_en": caption_en, "path": image_path}))
            i += 1
            continue

        if re.match(r"^表\d+(\.\d+)*\s+", stripped):
            caption_zh = stripped
            caption_en = ""
            i += 1
            if i < len(lines) and lines[i].strip().startswith("Tab "):
                caption_en = lines[i].strip()
                i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            rows, i = parse_markdown_table(lines, i)
            elements.append(("table", {"caption_zh": caption_zh, "caption_en": caption_en, "rows": rows}))
            continue

        if stripped.startswith("$$"):
            formula_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                formula_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                i += 1
            elements.append(("formula", "\n".join(formula_lines).strip()))
            continue

        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            elements.append(("bullet_list", items))
            continue

        para_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if not current:
                break
            if current.startswith(("# ", "## ", "### ", "![", "$$")):
                break
            if re.match(r"^表\d+(\.\d+)*\s+", current):
                break
            para_lines.append(current)
            i += 1
        elements.append(("paragraph", " ".join(para_lines)))

    return elements


def meta_get(meta, key, default=""):
    value = meta.get(key, default)
    return "" if value is None else str(value)


def split_paragraphs(text):
    return [part.strip() for part in re.split(r"\n\s*\n", text or "") if part.strip()]


def add_blank_lines(doc, count):
    for _ in range(count):
        doc.add_paragraph()


def build_cover(doc, meta):
    fields_top = [
        ("单位代码", meta_get(meta, "unit_code", "10359")),
        ("学    号", meta_get(meta, "student_id")),
        ("密  级", meta_get(meta, "security")),
        ("分类号", meta_get(meta, "classification")),
    ]
    for label, value in fields_top:
        add_text_paragraph(doc, f"{label}：{value}", FONT_HEI, FONT_EN, "小四", bold=True, exact_line=15)

    add_blank_lines(doc, 4)
    add_text_paragraph(doc, "Hefei University of Technology", FONT_SONG, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "硕士学位论文", FONT_HEI, FONT_EN, 48, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_text_paragraph(doc, "MASTER'S DISSERTATION", FONT_SONG, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_blank_lines(doc, 6)

    fields = [
        ("论文题目", meta_get(meta, "title_zh")),
        ("学位类别", meta_get(meta, "degree_category", "学历硕士")),
        ("专业名称", meta_get(meta, "major")),
        ("作者姓名", meta_get(meta, "author")),
        ("导师姓名", f"{meta_get(meta, 'advisor_name')}   {meta_get(meta, 'advisor_title')}".strip()),
        ("完成时间", meta_get(meta, "complete_time")),
    ]
    for label, value in fields:
        add_text_paragraph(doc, f"{label}：      {value}", FONT_SONG, FONT_EN, "三号", exact_line=20)

    add_blank_lines(doc, 2)
    add_text_paragraph(doc, "合   肥   工   业   大   学", FONT_HEI, FONT_EN, "二号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15, before=8, after=8)
    doc.add_page_break()


def build_chinese_title(doc, meta):
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "合   肥   工   业   大   学", FONT_HEI, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=18, before=8, after=8)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, f"{meta_get(meta, 'degree_category', '学历硕士')}学位论文", FONT_HEI, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15, before=8, after=8)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, meta_get(meta, "title_zh"), FONT_SONG, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15, before=8, after=8)
    add_blank_lines(doc, 5)

    fields = [
        ("作者姓名", meta_get(meta, "author")),
        ("指导教师", f"{meta_get(meta, 'advisor_name')}   {meta_get(meta, 'advisor_title')}".strip()),
        ("学科专业", meta_get(meta, "major")),
        ("研究方向", meta_get(meta, "research_direction")),
    ]
    for label, value in fields:
        add_text_paragraph(doc, f"{label}：    {value}", FONT_SONG, FONT_EN, "三号", exact_line=20, before=8, after=8)

    add_blank_lines(doc, 2)
    add_text_paragraph(doc, meta_get(meta, "complete_time"), FONT_SONG, FONT_EN, "小三", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15, before=8, after=8)
    doc.add_page_break()


def build_english_title(doc, meta):
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "A Dissertation Submitted for the Degree of Master", FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_blank_lines(doc, 2)
    add_text_paragraph(doc, meta_get(meta, "title_en"), FONT_SONG, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "By", FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_text_paragraph(doc, meta_get(meta, "author_en"), FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_blank_lines(doc, 8)
    add_text_paragraph(doc, "Hefei University of Technology", FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_text_paragraph(doc, "Hefei, Anhui, P.R.China", FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    add_text_paragraph(doc, f"{meta_get(meta, 'complete_month_en', 'Month')}, {meta_get(meta, 'complete_year_en', 'Year')}", FONT_SONG, FONT_EN, "三号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=15)
    doc.add_page_break()


def build_defense_page(doc, meta):
    committee = meta.get("defense_committee") or {}
    add_text_paragraph(doc, "合 肥 工 业 大 学", FONT_HEI, FONT_EN, "二号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=18, before=16, after=16)
    add_text_paragraph(doc, f"本论文经答辩委员会全体委员审查，确认符合合肥工业大学{meta_get(meta, 'degree_category', '学历硕士')}学位论文质量要求。", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "答辩委员会签名（工作单位、职称、姓名）", FONT_SONG, FONT_EN, "三号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=18, before=16, after=16)
    add_text_paragraph(doc, f"主席：{committee.get('chair', '')}", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
    members = committee.get("members") or []
    if isinstance(members, str):
        members = [members]
    if members:
        add_text_paragraph(doc, f"委员：{members[0]}", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
        for member in members[1:]:
            add_text_paragraph(doc, f"      {member}", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
    else:
        add_text_paragraph(doc, "委员：", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
    add_blank_lines(doc, 6)
    advisor = committee.get("advisor") or f"合肥工业大学，{meta_get(meta, 'advisor_title')}，{meta_get(meta, 'advisor_name')}"
    add_text_paragraph(doc, f"导师：{advisor}", FONT_SONG, FONT_EN, "三号", first_indent=2, exact_line=18, before=16, after=16)
    doc.add_page_break()


def build_declaration(doc, meta):
    add_text_paragraph(doc, "学位论文独创性声明", FONT_HEI, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=18, before=8, after=24)
    text1 = "本人郑重声明：所呈交的学位论文是本人在导师指导下进行独立研究工作所取得的成果。据我所知，除了文中特别加以标注和致谢的内容外，论文中不包含其他人已经发表或撰写过的研究成果，也不包含为获得  合肥工业大学   或其他教育机构的学位或证书而使用过的材料。对本文成果做出贡献的个人和集体，本人已在论文中作了明确的说明，并表示谢意。"
    text2 = "学位论文中表达的观点纯属作者本人观点，与合肥工业大学无关。"
    add_text_paragraph(doc, text1, FONT_SONG, FONT_EN, "小四", first_indent=2, exact_line=22)
    add_text_paragraph(doc, text2, FONT_SONG, FONT_EN, "小四", first_indent=2, exact_line=22)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "学位论文作者签名：　　　　　　　　　签名日期：　　　年　　月　　日", FONT_SONG, FONT_EN, "小四", first_indent=2, exact_line=22)
    add_blank_lines(doc, 2)

    add_text_paragraph(doc, "学位论文版权使用授权书", FONT_HEI, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=18, before=8, after=24)
    text3 = "本学位论文作者完全了解  合肥工业大学  有关保留、使用学位论文的规定，即：除保密期内的涉密学位论文外，学校有权保存并向国家有关部门或机构送交论文的复印件和电子光盘，允许论文被查阅或借阅。本人授权  合肥工业大学  可以将本学位论文的全部或部分内容编入有关数据库，允许采用影印、缩印或扫描等复制手段保存、汇编学位论文。"
    add_text_paragraph(doc, text3, FONT_SONG, FONT_EN, "小四", first_indent=2, exact_line=22)
    add_text_paragraph(doc, "（保密的学位论文在解密后适用本授权书）", FONT_SONG, FONT_EN, "小四", first_indent=2, exact_line=22)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, "学位论文作者签名：", FONT_SONG, FONT_EN, "小四", exact_line=22, before=8, after=8)
    add_text_paragraph(doc, "指导教师签名：", FONT_SONG, FONT_EN, "小四", exact_line=22, before=8, after=8)
    add_text_paragraph(doc, "签名日期：　　　年   月   日", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_text_paragraph(doc, "签名日期：　　　年   月   日", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_blank_lines(doc, 2)

    dest = meta.get("graduation_destination") or {}
    add_text_paragraph(doc, "论文作者毕业去向", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_blank_lines(doc, 1)
    add_text_paragraph(doc, f"工作单位：{dest.get('work_unit', '')}", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_text_paragraph(doc, f"联系电话：{dest.get('telephone', '')}", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_text_paragraph(doc, f"E-mail：{dest.get('email', '')}", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_text_paragraph(doc, f"通讯地址：{dest.get('address', '')}", FONT_SONG, FONT_EN, "小四", exact_line=22)
    add_text_paragraph(doc, f"邮政编码：{dest.get('postcode', '')}", FONT_SONG, FONT_EN, "小四", exact_line=22)


def build_section_title(doc, title):
    add_text_paragraph(doc, title, FONT_SONG, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=8, after=24)


def add_body_paragraph(doc, text, first_indent=2):
    return add_text_paragraph(doc, text, FONT_SONG, FONT_EN, "小四", first_indent=first_indent, exact_line=20)


def add_keyword_line(doc, prefix, words, english=False):
    if not words:
        return
    sep = "; " if english else "；"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing_exact(p, 20)
    set_spacing(p, before=16, after=0)
    run = p.add_run(prefix)
    set_run_font(run, FONT_EN if english else FONT_SONG, FONT_EN, "小四", bold=True)
    run = p.add_run(sep.join(str(w) for w in words))
    set_run_font(run, FONT_EN if english else FONT_SONG, FONT_EN, "小四")


def build_acknowledgments(doc, meta):
    build_section_title(doc, "致  谢")
    for para in split_paragraphs(meta_get(meta, "acknowledgments")):
        add_body_paragraph(doc, para)
    doc.add_page_break()


def build_abstract_zh(doc, meta):
    build_section_title(doc, "摘  要")
    for para in split_paragraphs(meta_get(meta, "abstract_zh")):
        add_body_paragraph(doc, para)
    add_keyword_line(doc, "关键词：", meta.get("keywords_zh") or [], english=False)
    doc.add_page_break()


def build_abstract_en(doc, meta):
    build_section_title(doc, "ABSTRACT")
    for para in split_paragraphs(meta_get(meta, "abstract_en")):
        add_text_paragraph(doc, para, FONT_EN, FONT_EN, "小四", first_indent=1, exact_line=20)
    add_keyword_line(doc, "KEYWORDS: ", meta.get("keywords_en") or [], english=True)
    doc.add_page_break()


def build_toc_placeholder(doc, meta):
    build_section_title(doc, "目  录")
    add_text_paragraph(doc, "请在 Word/WPS 中插入或更新自动目录（显示到第三级标题）。", FONT_SONG, FONT_EN, "小四", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=20)
    doc.add_page_break()

    if meta.get("include_figure_list", True):
        build_section_title(doc, "插图清单")
        add_text_paragraph(doc, "请在 Word/WPS 中根据图题更新插图清单。", FONT_SONG, FONT_EN, "小四", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=20)
        doc.add_page_break()

    if meta.get("include_table_list", True):
        build_section_title(doc, "表格清单")
        add_text_paragraph(doc, "请在 Word/WPS 中根据表题更新表格清单。", FONT_SONG, FONT_EN, "小四", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=20)
        doc.add_page_break()

    symbols = meta_get(meta, "symbols").strip()
    if symbols:
        build_section_title(doc, "主要符号注释表")
        for para in split_paragraphs(symbols):
            add_body_paragraph(doc, para, first_indent=0)
        doc.add_page_break()


def build_body(doc, elements, source_dir):
    first_chapter = True
    for element_type, data in elements:
        if element_type == "chapter":
            if not first_chapter:
                doc.add_page_break()
            first_chapter = False
            p = add_text_paragraph(doc, data, FONT_HEI, FONT_EN, "三号", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=12, after=12, style="Heading 1")
            for run in p.runs:
                set_run_font(run, FONT_HEI, FONT_EN, "三号", bold=True)

        elif element_type == "section":
            p = add_text_paragraph(doc, data, FONT_HEI, FONT_EN, "小四", bold=True, exact_line=None, before=6, after=6, style="Heading 2")
            for run in p.runs:
                set_run_font(run, FONT_HEI, FONT_EN, "小四", bold=True)

        elif element_type == "subsection":
            p = add_text_paragraph(doc, data, FONT_SONG, FONT_EN, "小四", bold=True, exact_line=None, before=6, after=6, style="Heading 3")
            for run in p.runs:
                set_run_font(run, FONT_SONG, FONT_EN, "小四", bold=True)

        elif element_type == "paragraph":
            add_body_paragraph(doc, data)

        elif element_type == "bullet_list":
            for item in data:
                p = add_text_paragraph(doc, item, FONT_SONG, FONT_EN, "小四", exact_line=20)
                p.paragraph_format.left_indent = Pt(24)
                p.paragraph_format.first_line_indent = Pt(-12)

        elif element_type == "figure":
            add_figure(doc, data, source_dir)

        elif element_type == "table":
            add_table(doc, data)

        elif element_type == "formula":
            add_text_paragraph(doc, data, FONT_EN, FONT_EN, "小四", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=20, before=6, after=6)


def add_figure(doc, data, source_dir):
    image_path = Path(data["path"])
    if not image_path.is_absolute():
        image_path = Path(source_dir) / image_path

    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(12))
    else:
        p = add_text_paragraph(doc, f"[图片缺失：{data['path']}]", FONT_SONG, FONT_EN, "小四", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=20)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    add_text_paragraph(doc, data["caption_zh"], FONT_SONG, FONT_EN, "五号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=6, after=0)
    if data.get("caption_en"):
        add_text_paragraph(doc, data["caption_en"], FONT_SONG, FONT_EN, "五号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=0, after=6)


def add_table(doc, data):
    add_text_paragraph(doc, data["caption_zh"], FONT_SONG, FONT_EN, "五号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=6, after=0)
    if data.get("caption_en"):
        add_text_paragraph(doc, data["caption_en"], FONT_SONG, FONT_EN, "五号", align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=0, after=6)

    rows = data.get("rows") or []
    if not rows:
        return

    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(ncols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            clear_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_single_line(p)
            set_spacing(p, before=3, after=3)
            run = p.add_run(text)
            set_run_font(run, FONT_SONG, FONT_EN, "五号", bold=(r_idx == 0))

    apply_three_line_table(table)


def set_border(element, edge, value="single", size="8", color="000000"):
    tag = qn(f"w:{edge}")
    border = element.find(tag)
    if border is None:
        border = OxmlElement(f"w:{edge}")
        element.append(border)
    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def apply_three_line_table(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    set_border(borders, "top", size="12")
    set_border(borders, "bottom", size="12")
    for edge in ("left", "right", "insideV", "insideH"):
        set_border(borders, edge, value="none", size="0")

    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = tc_pr.find(qn("w:tcBorders"))
            if cell_borders is None:
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
            set_border(cell_borders, "bottom", size="6")


def build_references(doc, meta):
    doc.add_page_break()
    p = add_text_paragraph(doc, "参考文献", FONT_SONG, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=8, after=24, style="Heading 1")
    for run in p.runs:
        set_run_font(run, FONT_SONG, FONT_EN, "小二", bold=True)

    for ref in meta.get("references") or []:
        p = add_text_paragraph(doc, str(ref), FONT_SONG, FONT_EN, "五号", exact_line=20)
        set_hanging_indent(p, chars=2.5, size="五号")


def build_appendix(doc, meta):
    text = meta_get(meta, "appendix").strip()
    if not text:
        return
    doc.add_page_break()
    p = add_text_paragraph(doc, "附录", FONT_SONG, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=8, after=24, style="Heading 1")
    for run in p.runs:
        set_run_font(run, FONT_SONG, FONT_EN, "小二", bold=True)
    for para in split_paragraphs(text):
        add_body_paragraph(doc, para)


def build_academic_achievements(doc, meta):
    text = meta_get(meta, "academic_achievements").strip()
    if not text:
        return
    doc.add_page_break()
    p = add_text_paragraph(doc, "攻读硕士学位期间的学术活动及成果情况", FONT_SONG, FONT_EN, "小二", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, exact_line=None, before=8, after=24, style="Heading 1")
    for run in p.runs:
        set_run_font(run, FONT_SONG, FONT_EN, "小二", bold=True)
    for para in split_paragraphs(text):
        add_body_paragraph(doc, para, first_indent=0)


def configure_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(FONT_SIZE["小四"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.color.rgb = RGBColor(0, 0, 0)


def build_document(meta, body, source_dir, output_path):
    doc = Document()
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        pass
    configure_base_styles(doc)

    first_section = doc.sections[0]
    set_section_page(first_section)
    build_cover(doc, meta)
    build_chinese_title(doc, meta)
    build_english_title(doc, meta)
    build_defense_page(doc, meta)
    build_declaration(doc, meta)

    roman_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page(roman_section)
    unlink_headers_footers(roman_section)
    set_page_numbering(roman_section, fmt="upperRoman", start=1)
    set_footer_page_number(roman_section)
    build_acknowledgments(doc, meta)
    build_abstract_zh(doc, meta)
    build_abstract_en(doc, meta)
    build_toc_placeholder(doc, meta)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page(body_section)
    unlink_headers_footers(body_section)
    set_page_numbering(body_section, fmt="decimal", start=1)
    set_footer_page_number(body_section)
    set_body_headers(body_section)

    elements = parse_body(body)
    build_body(doc, elements, source_dir)
    build_references(doc, meta)
    build_appendix(doc, meta)
    build_academic_achievements(doc, meta)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python build_thesis.py input.md [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else input_path.with_suffix(".docx")
    meta, body = parse_source(input_path)
    result = build_document(meta, body, input_path.parent, output_path)
    print(f"Generated: {result}")
    print("Note: update TOC/list/page-number fields in Word/WPS before final submission.")


if __name__ == "__main__":
    main()
